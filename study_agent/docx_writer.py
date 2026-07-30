"""LLMが返すMarkdownを .docx に変換する簡易ライター (python-docx)。

対応: # 見出し1〜3 / - 箇条書き / 1. 番号付き / **太字**(行内) / 普通の段落。
"""
import re
from pathlib import Path

import docx
from docx.shared import Pt, RGBColor

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str):
    """**bold** を解釈しながらランを追加。"""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(md: str, out_path: str, title: str | None = None,
                     banner: str | None = None) -> str:
    doc = docx.Document()

    if title:
        doc.add_heading(title, level=0)

    if banner:
        p = doc.add_paragraph()
        run = p.add_run(banner)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
